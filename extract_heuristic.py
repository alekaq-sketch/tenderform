"""
Извлечение позиций БЕЗ обращения к LLM - бесплатно, без API-ключей.

Идея: вместо того, чтобы полагаться на номер колонки (что и ломалось на
разных шаблонах), ищем в каждой таблице строку-заголовок по ключевым словам
(на русском и казахском), а дальше читаем данные под найденными колонками.
Это работает на большинстве реальных документов, потому что слова вроде
"наименование/атауы", "кол-во/саны", "цена/баға" почти всегда присутствуют -
даже если порядок и количество колонок отличаются.

Ограничение: если в документе вообще нет таблицы (только текст в прозе:
"поставить 720 тонн цианида натрия по цене...") - это уже не по силам
ни этому скрипту, ни regex в принципе, тут поможет либо ручной ввод,
либо LLM-шаг (extract_with_llm.py) для конкретно этого документа.
"""
import re

HEADER_KEYWORDS = {
    "name": ["наименован", "нименован", "атауы", "товар"],
    "unit": ["ед.изм", "ед. изм", "единиц", "өлш"],
    "qty": ["кол-во", "количество", "саны"],
    "price": ["цена", "баға", "стоимост"],
}
STOP_ROW_KEYWORDS = ["итого", "барлығы", "всего", "жиынтығ"]


def _norm(s):
    return str(s or "").strip().lower()


def _to_number(s):
    if s is None:
        return None
    s = str(s).strip().replace(" ", "").replace(",", ".")
    s = re.sub(r"[^\d.]", "", s)
    if not s or s == ".":
        return None
    try:
        return float(s) if "." in s else int(s)
    except ValueError:
        return None


def _find_header_row(table):
    """Возвращает (индекс строки, {field: col_index}) или (None, {})."""
    if not table or max(len(row) for row in table) < 2:
        return None, {}  # однострочная "таблица" - это абзац текста, не таблица позиций

    best_row, best_map, best_score = None, {}, 0
    for ri, row in enumerate(table):
        col_map = {}
        for ci, cell in enumerate(row):
            cell_norm = _norm(cell)
            if len(cell_norm) > 120:
                continue  # заголовки колонок - короткие подписи (даже двуязычные), не абзацы прозы
            for field, keywords in HEADER_KEYWORDS.items():
                if field not in col_map and any(k in cell_norm for k in keywords):
                    col_map[field] = ci
        score = len(col_map)
        if score > best_score and "name" in col_map:  # наименование обязательно
            best_row, best_map, best_score = ri, col_map, score
    if best_score >= 2:  # нашли минимум 2 поля - считаем, что это заголовок
        return best_row, best_map
    return None, {}


def parse_table(table):
    """table - список строк, каждая строка - список значений ячеек."""
    header_idx, col_map = _find_header_row(table)
    if header_idx is None:
        return []

    items = []
    for row in table[header_idx + 1:]:
        name = row[col_map["name"]] if col_map.get("name") is not None and col_map["name"] < len(row) else None
        if not name or not str(name).strip():
            continue
        if any(k in _norm(name) for k in STOP_ROW_KEYWORDS):
            break

        def get(field):
            idx = col_map.get(field)
            if idx is None or idx >= len(row):
                return None
            return row[idx]

        items.append({
            "name": str(name).strip(),
            "unit": (str(get("unit")).strip() if get("unit") else None),
            "qty": _to_number(get("qty")),
            "unit_price": _to_number(get("price")),
        })
    return items


def docx_tables(path):
    from docx import Document
    doc = Document(path)

    def walk(tables):
        for table in tables:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cells.append(cell.text.strip())
                    if cell.tables:
                        yield from walk(cell.tables)
                rows.append(cells)
            yield rows

    return list(walk(doc.tables))


def xlsx_tables(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    tables = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(v is not None for v in row):
                rows.append(list(row))
        if rows:
            tables.append(rows)
    return tables


def pdf_tables(path):
    import pdfplumber
    tables = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tables.extend(page.extract_tables())
    return tables


def extract_items(path: str) -> list[dict]:
    lower_path = path.lower()
    if lower_path.endswith(".docx"):
        tables = docx_tables(path)
    elif lower_path.endswith(".xlsx"):
        tables = xlsx_tables(path)
    elif lower_path.endswith(".pdf"):
        tables = pdf_tables(path)
    else:
        raise ValueError(f"Неизвестный формат: {path}")

    all_items = []
    for table in tables:
        all_items.extend(parse_table(table))
    return all_items


if __name__ == "__main__":
    import sys
    import json
    print(json.dumps(extract_items(sys.argv[1]), ensure_ascii=False, indent=2))
