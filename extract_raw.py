"""
Шаг 1 пайплайна - выгрузка ВСЕГО текста и ВСЕХ таблиц из документа, включая
вложенные таблицы в docx. Это не пытается понять, что означает какая колонка -
просто гарантированно достаёт всё содержимое в текстовом виде, независимо от
того, насколько структура документа отличается от предыдущего лота.

Именно этот шаг надёжен всегда. Смысловое сопоставление полей - шаг 2,
в extract_with_llm.py.
"""


def _walk_docx_tables(doc, tables):
    """Рекурсивно собирает все таблицы, включая вложенные в ячейки."""
    for table in tables:
        rows_text = []
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                # текст самой ячейки
                cells_text.append(cell.text.strip())
                # + рекурсивно вложенные таблицы внутри этой ячейки
                if cell.tables:
                    yield from _walk_docx_tables(doc, cell.tables)
            rows_text.append(cells_text)
        yield rows_text


def extract_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)

    chunks = []
    chunks.append("=== ТЕКСТ ДОКУМЕНТА ===")
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text.strip())

    chunks.append("\n=== ТАБЛИЦЫ (включая вложенные) ===")
    for i, table_rows in enumerate(_walk_docx_tables(doc, doc.tables)):
        chunks.append(f"\n--- таблица {i} ---")
        for row in table_rows:
            chunks.append(" | ".join(row))

    return "\n".join(chunks)


def extract_xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    chunks = []
    for ws in wb.worksheets:
        chunks.append(f"\n=== ЛИСТ: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            if any(v is not None for v in row):
                chunks.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(chunks)


def extract_pdf(path: str) -> str:
    import pdfplumber
    chunks = []
    with pdfplumber.open(path) as pdf:
        for pi, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            chunks.append(f"\n=== страница {pi + 1} ===\n{text}")
            for table in page.extract_tables():
                chunks.append("--- таблица ---")
                for row in table:
                    chunks.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(chunks)


def extract_any(path: str) -> str:
    lower_path = path.lower()
    if lower_path.endswith(".docx"):
        return extract_docx(path)
    if lower_path.endswith(".xlsx"):
        return extract_xlsx(path)
    if lower_path.endswith(".pdf"):
        return extract_pdf(path)
    raise ValueError(f"Неизвестный формат: {path}")


if __name__ == "__main__":
    import sys
    print(extract_any(sys.argv[1]))
